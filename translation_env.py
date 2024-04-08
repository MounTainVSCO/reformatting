
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

# Paragraph-Level Formatting: 

# Applies to the entire paragraph and affects aspects like alignment, line spacing, and indentation. 
# It does not directly change the formatting of text within the paragraph's 
# runs, such as font size, bold, italic, etc., unless those runs are explicitly formatted to "inherit" from the paragraph.


# Run-Level Formatting: 

# Applies to specific segments of text within a paragraph, controlling font type, size, bold, italic, underline, and color. 
# Run-level formatting takes precedence over paragraph-level formatting for these attributes.


# translation of text with context: get original paragraph -> get paragraph styles and runs! -> start translation of original text 
    
def paragraph_and_runs_formatting(file_path):
    # Get a motha fuckin paragraph
    doc = Document(file_path)
    text_and_formatting = {}

    for para in doc.paragraphs:
        para_text = para.text

        format = para.paragraph_format
        para_format = {
            "alignment": para.alignment,
            "style": para.style,
            "left_indent": format.left_indent,
            "right_indent":format.right_indent,
            "first_line_indent":format.first_line_indent
        }
        runs = []
        for run in para.runs:
            if (not run.text or run.text == " "): continue # new line
            run_details = {
                "text": run.text,
                "bold": run.bold,
                "italic": run.italic,
                "underline": run.underline,
                "font": run.font.name,
                "size": run.font.size,
            }
            runs.append(run_details)
        text_and_formatting[para_text] = [para_format, runs]
    return text_and_formatting


def reformatted_text(paragraph_runs):

    paragraph = doc.add_paragraph()
    for text, paragraph_and_runs_styles in paragraph_runs.items():
        
        paragraph_styles = paragraph_and_runs_styles[0]

        all_runs_in_paragraph = paragraph_and_runs_styles[1]
        for run in all_runs_in_paragraph:
            runs = paragraph.add_run(f"{run['text']} ")
            if run['bold'] is not None:
                runs.bold = run['bold']
            if run['italic'] is not None:
                runs.italic = run['italic']
            if run['underline'] is not None:
                runs.underline = run['underline']
            if run['font'] is not None:
                runs.font.name = run['font']
            if run['size'] is not None:
                runs.font.size = Pt(run['size'] / 12700)
            
        if paragraph_styles['alignment'] ==  WD_PARAGRAPH_ALIGNMENT.RIGHT: paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
        elif paragraph_styles['alignment'] ==  WD_PARAGRAPH_ALIGNMENT.LEFT: paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        elif paragraph_styles['alignment'] ==  WD_PARAGRAPH_ALIGNMENT.CENTER: paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        if paragraph_styles['style']: paragraph.style = paragraph_styles['style']
        if paragraph_styles['alignment']: paragraph.alignment = paragraph_styles['alignment']
        if paragraph_styles['left_indent']: paragraph.left_indent = paragraph_styles['left_indent']
        if paragraph_styles['right_indent']: paragraph.right_indent = paragraph_styles['right_indent']
        if paragraph_styles['first_line_indent']: paragraph.first_line_indent = paragraph_styles['first_line_indent']

        doc.save(output_file)
    doc.save(output_file)

if (__name__ == "__main__"):
    doc = Document()

    input_file ="input_file/reformatting env(3).docx"
    output_file = "output_file/out-reformatting env(3).docx"
    text_formatting = paragraph_and_runs_formatting(input_file)
    reformatted_text(text_formatting)
